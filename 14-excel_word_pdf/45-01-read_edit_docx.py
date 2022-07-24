#!/usr/bin/env python
# pip install python-docx
import docx

# Document contains paragraphs
d = docx.Document("demo.docx")
print(d.paragraphs)

print(d.paragraphs[0].text)
print(d.paragraphs[1].text)

# Paragraph contains runs
# Each run contain the text with the same style
p = d.paragraphs[1]
print(p.runs)
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)

print(p.runs[1].bold)
print(p.runs[3].italic)

# Make changes
p.runs[3].underline = True
p.runs[3].text = "italic and underline"

# Save to file
d.save("demo2.docx")

p.style = "Title"
d.save("demo3.docx")
